# customers.py

import streamlit as st
import pandas as pd
from datetime import datetime, date
import time
import io
from database import query_df, execute_sql
from utils import get_user_map, clear_user_cache
from config import CUSTOMER_LEVELS, CUSTOMER_SOURCES
from validators import validate_phone, validate_required


# 导入 Word 生成库
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@st.cache_data(ttl=300, show_spinner="正在加载客户数据...")
def load_customers_data(uid: str, is_boss: bool):
    """加载客户数据（带缓存）"""
    if is_boss:
        df = query_df("SELECT id, name, phone, company, level, source, owner_id, last_follow, created_at FROM customers ORDER BY created_at DESC")
    else:
        df = query_df("SELECT id, name, phone, company, level, source, owner_id, last_follow, created_at FROM customers WHERE owner_id = ? ORDER BY created_at DESC", (uid,))
    return df


@st.cache_data(ttl=300, show_spinner="正在加载客户列表...")
def load_customers_list(uid: str, is_boss: bool):
    """加载客户选择列表（带缓存）"""
    if is_boss:
        df = query_df("SELECT id, name, company, owner_id FROM customers ORDER BY name")
    else:
        df = query_df("SELECT id, name, company, owner_id FROM customers WHERE owner_id = ? ORDER BY name", (uid,))
    return df


def filter_customers(df, search_term, level_filter, source_filter):
    """筛选客户数据"""
    filtered = df.copy()
    
    # 搜索筛选
    if search_term:
        search_term = search_term.lower()
        mask = (
            filtered['name'].str.lower().str.contains(search_term, na=False) |
            filtered['company'].str.lower().str.contains(search_term, na=False) |
            filtered['phone'].str.lower().str.contains(search_term, na=False)
        )
        filtered = filtered[mask]
    
    # 等级筛选
    if level_filter != "全部":
        filtered = filtered[filtered['level'] == level_filter]
    
    # 来源筛选
    if source_filter != "全部":
        filtered = filtered[filtered['source'] == source_filter]
    
    return filtered


def paginate_dataframe(df, page_size=20):
    """分页数据"""
    total_pages = max(1, (len(df) - 1) // page_size + 1)
    return total_pages


def show_customers(uid: str, is_boss: bool):
    """客户管理（优化版本）"""
    st.title("👥 客户资源库")
    
    # 初始化会话状态
    if 'customer_page' not in st.session_state:
        st.session_state.customer_page = 1
    if 'search_term' not in st.session_state:
        st.session_state.search_term = ""
    if 'level_filter' not in st.session_state:
        st.session_state.level_filter = "全部"
    if 'source_filter' not in st.session_state:
        st.session_state.source_filter = "全部"
    
    user_map = get_user_map()
    
    # 创建标签页
    tab1, tab2, tab3, tab4 = st.tabs(["📋 客户列表", "➕ 新增客户", "✏️ 编辑客户", "📝 跟进日志"])
    
    # ---------- 标签页1：客户列表（优化版本） ----------
    with tab1:
        # 筛选区域（放在顶部）
        st.subheader("🔍 筛选和搜索")
        col1, col2, col3 = st.columns(3)
        with col1:
            search_term = st.text_input("搜索客户（姓名/公司/手机）", 
                                       value=st.session_state.search_term,
                                       placeholder="输入关键词搜索...")
        with col2:
            level_options = ["全部"] + CUSTOMER_LEVELS
            try:
                level_index = level_options.index(st.session_state.level_filter)
            except ValueError:
                level_index = 0
            level_filter = st.selectbox("客户等级", 
                                      level_options,
                                      index=level_index)
        with col3:
            source_options = ["全部"] + CUSTOMER_SOURCES
            try:
                source_index = source_options.index(st.session_state.source_filter)
            except ValueError:
                source_index = 0
            source_filter = st.selectbox("客户来源", 
                                       source_options,
                                       index=source_index)
        
        # 更新会话状态
        st.session_state.search_term = search_term
        st.session_state.level_filter = level_filter
        st.session_state.source_filter = source_filter
        
        # 加载并筛选数据
        df_cust = load_customers_data(uid, is_boss)
        if not df_cust.empty:
            df_cust = filter_customers(df_cust, search_term, level_filter, source_filter)
            df_cust["owner_name"] = df_cust["owner_id"].map(user_map).fillna(df_cust["owner_id"])
            
            # 分页显示
            page_size = 20
            total_pages = paginate_dataframe(df_cust, page_size)
            
            # 分页控件
            if total_pages > 1:
                col_p1, col_p2, col_p3 = st.columns([1, 2, 1])
                with col_p1:
                    if st.button("⬅️ 上一页", disabled=st.session_state.customer_page <= 1):
                        st.session_state.customer_page = max(1, st.session_state.customer_page - 1)
                        st.rerun()
                with col_p2:
                    st.info(f"第 {st.session_state.customer_page} / {total_pages} 页，共 {len(df_cust)} 条记录")
                with col_p3:
                    if st.button("➡️ 下一页", disabled=st.session_state.customer_page >= total_pages):
                        st.session_state.customer_page = min(total_pages, st.session_state.customer_page + 1)
                        st.rerun()
            
            # 计算当前页数据
            start_idx = (st.session_state.customer_page - 1) * page_size
            end_idx = start_idx + page_size
            df_display = df_cust.iloc[start_idx:end_idx].copy()
            
            # 格式化日期列
            display_cols = ["name", "phone", "company", "level", "source", "owner_name", "last_follow", "created_at"]
            display_cols = [c for c in display_cols if c in df_display.columns]
            
            for col in ["last_follow", "created_at"]:
                if col in df_display.columns:
                    df_display[col] = pd.to_datetime(df_display[col]).dt.date
            
            # 显示数据表格
            column_config = {
                "name": "联系人",
                "phone": "手机号",
                "company": "公司名称",
                "level": "客户等级",
                "source": "来源",
                "owner_name": "负责人",
                "last_follow": st.column_config.DateColumn("最后跟进", format="YYYY-MM-DD"),
                "created_at": st.column_config.DateColumn("创建时间", format="YYYY-MM-DD")
            }
            
            st.dataframe(df_display[display_cols], 
                        use_container_width=True, 
                        hide_index=True, 
                        column_config=column_config)
        else:
            st.info("暂无客户数据")
    
    # ---------- 标签页2：新增客户 ----------
    with tab2:
        with st.form("new_customer", clear_on_submit=True):
            st.subheader("新增客户信息")
            name = st.text_input("联系人 *")
            phone = st.text_input("手机号", placeholder="例如：13800138000")
            company = st.text_input("公司名称")
            level = st.selectbox("客户等级", CUSTOMER_LEVELS)
            source = st.selectbox("来源", CUSTOMER_SOURCES)
            
            if st.form_submit_button("确认录入", type="primary"):
                # 验证输入
                is_valid, error_msg = validate_required(name, "联系人")
                if not is_valid:
                    st.error(error_msg)
                else:
                    is_valid, error_msg = validate_phone(phone)
                    if not is_valid:
                        st.error(error_msg)
                    else:
                        with st.spinner("正在保存客户..."):
                            today = date.today()
                            sql = """
                                INSERT INTO customers
                                (name, phone, company, level, source, owner_id, last_follow, created_at)
                                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            """
                            params = (name.strip(), phone.strip() if phone else None, company, level, source, uid, today, today)
                            try:
                                execute_sql(sql, params)
                                st.success("客户入库成功")
                                clear_user_cache()
                                # 清除相关缓存
                                load_customers_data.clear()
                                load_customers_list.clear()
                                time.sleep(0.5)
                                st.rerun()
                            except Exception as e:
                                st.error(f"保存失败: {e}")
    
    # ---------- 标签页3：编辑客户 ----------
    with tab3:
        cust_list = load_customers_list(uid, is_boss)
        if cust_list.empty:
            st.info("暂无您负责的客户")
        else:
            cust_options = {f"{row['id']} - {row['name']} ({row['company']})": row['id'] for _, row in cust_list.iterrows()}
            selected_label = st.selectbox("选择要编辑的客户", list(cust_options.keys()))
            selected_id = cust_options[selected_label]
            
            # 获取客户详情
            cust_info = query_df("SELECT * FROM customers WHERE id = ?", (selected_id,))
            if cust_info.empty:
                st.error("客户不存在")
            else:
                row = cust_info.iloc[0]
                with st.form("edit_customer"):
                    st.subheader("编辑客户信息")
                    new_name = st.text_input("联系人", value=row["name"])
                    new_phone = st.text_input("手机号", value=row["phone"] if pd.notna(row["phone"]) else "")
                    new_company = st.text_input("公司名称", value=row["company"] if pd.notna(row["company"]) else "")
                    new_level = st.selectbox("客户等级", CUSTOMER_LEVELS,
                                          index=CUSTOMER_LEVELS.index(row["level"]) if row["level"] in CUSTOMER_LEVELS else 0)
                    new_source = st.selectbox("来源", CUSTOMER_SOURCES,
                                           index=CUSTOMER_SOURCES.index(row["source"]) if row["source"] in CUSTOMER_SOURCES else 0)
                    new_last_follow = st.date_input("最后跟进日期",
                                                  value=pd.to_datetime(row["last_follow"]).date() if pd.notna(row["last_follow"]) else date.today())
                    
                    if is_boss:
                        all_users = query_df("SELECT username, name FROM users ORDER BY name")
                        user_options = {f"{u['name']} ({u['username']})": u['username'] for _, u in all_users.iterrows()}
                        current_owner_label = next((label for label, un in user_options.items() if un == row["owner_id"]),
                                                 f"{row['owner_id']} - {row['owner_id']}")
                        selected_owner_label = st.selectbox("负责人", list(user_options.keys()),
                                                          index=list(user_options.keys()).index(current_owner_label) if current_owner_label in user_options else 0)
                        new_owner = user_options[selected_owner_label]
                    else:
                        new_owner = row["owner_id"]
                    
                    if st.form_submit_button("更新客户", type="primary"):
                        # 验证输入
                        is_valid, error_msg = validate_required(new_name, "联系人")
                        if not is_valid:
                            st.error(error_msg)
                        else:
                            is_valid, error_msg = validate_phone(new_phone)
                            if not is_valid:
                                st.error(error_msg)
                            else:
                                with st.spinner("正在更新客户..."):
                                    sql = """
                                        UPDATE customers
                                        SET name=?, phone=?, company=?, level=?, source=?, last_follow=?, owner_id=?
                                        WHERE id=?
                                    """
                                    params = (new_name.strip(), new_phone.strip() if new_phone else None, new_company, new_level, new_source, new_last_follow, new_owner, selected_id)
                                    try:
                                        execute_sql(sql, params)
                                        st.success("客户信息已更新")
                                        clear_user_cache()
                                        # 清除相关缓存
                                        load_customers_data.clear()
                                        load_customers_list.clear()
                                        time.sleep(0.5)
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"更新失败: {e}")
                
                st.divider()
                
                # 删除客户功能
                biz_check = query_df("SELECT id FROM business WHERE cust_id = ?", (selected_id,))
                if not biz_check.empty:
                    st.error("该客户存在关联商机，无法删除。请先删除商机。")
                else:
                    with st.popover("🗑️ 删除客户"):
                        st.warning("删除客户将同时删除所有跟进记录，此操作不可逆。")
                        if st.button("确认删除", key=f"confirm_del_cust_{selected_id}", type="primary"):
                            with st.spinner("正在删除客户及相关记录..."):
                                try:
                                    execute_sql("DELETE FROM follow_logs WHERE ref_type='customer' AND ref_id=?", (selected_id,))
                                    execute_sql("DELETE FROM customers WHERE id=?", (selected_id,))
                                    st.success("客户已删除")
                                    clear_user_cache()
                                    # 清除相关缓存
                                    load_customers_data.clear()
                                    load_customers_list.clear()
                                    time.sleep(0.5)
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"删除失败: {e}")
    
    # ---------- 标签页4：跟进日志 ----------
    with tab4:
        st.subheader("客户跟进历史")
        
        cust_list = load_customers_list(uid, is_boss)
        if cust_list.empty:
            st.info("暂无客户")
            return
        
        # 导出 Excel 功能
        st.markdown("### 📎 导出跟进日志（Excel）")
        export_type = st.radio("导出范围", ["当前选中的客户", "所有客户（按权限）"], horizontal=True, key="export_type")
        
        if export_type == "当前选中的客户":
            cust_options = {f"{row['id']} - {row['name']} ({row['company']})": row['id'] for _, row in cust_list.iterrows()}
            selected_label = st.selectbox("选择客户", list(cust_options.keys()), key="export_cust_select")
            selected_cust_id = cust_options[selected_label]
            
            if st.button("导出当前客户跟进日志"):
                with st.spinner("正在生成导出文件..."):
                    logs = query_df("""
                        SELECT fl.*, u.name as user_name, coalesce(fl.log_time, fl.created_at) as display_time
                        FROM follow_logs fl
                        LEFT JOIN users u ON fl.user_id = u.username
                        WHERE fl.ref_type='customer' AND fl.ref_id = ?
                        ORDER BY display_time DESC
                    """, (selected_cust_id,))
                    if not logs.empty:
                        cust_info = cust_list[cust_list['id'] == selected_cust_id].iloc[0]
                        logs.insert(0, '客户名称', cust_info['name'])
                        logs.insert(1, '公司名称', cust_info['company'])
                        df_export = logs[['客户名称', '公司名称', 'user_name', 'display_time', 'subject', 'content', 'participants', 'location', 'next_plan']]
                        df_export.columns = ['客户名称', '公司名称', '跟进人', '跟进时间', '主题', '内容', '相关人员', '地点', '下一步计划']
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            df_export.to_excel(writer, sheet_name='跟进日志', index=False)
                        output.seek(0)
                        st.download_button(
                            label="📥 下载 Excel",
                            data=output,
                            file_name=f"跟进日志_{cust_info['name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    else:
                        st.warning("该客户暂无跟进记录")
        else:
            if st.button("导出所有客户跟进日志"):
                with st.spinner("正在生成导出文件..."):
                    cust_ids = cust_list['id'].tolist()
                    if not cust_ids:
                        st.warning("没有可导出的客户")
                    else:
                        placeholders = ','.join(['?']*len(cust_ids))
                        logs = query_df(f"""
                            SELECT fl.*, u.name as user_name, c.name as customer_name, c.company as customer_company,
                                   coalesce(fl.log_time, fl.created_at) as display_time
                            FROM follow_logs fl
                            LEFT JOIN users u ON fl.user_id = u.username
                            LEFT JOIN customers c ON fl.ref_id = c.id
                            WHERE fl.ref_type='customer' AND fl.ref_id IN ({placeholders})
                            ORDER BY c.name, display_time DESC
                        """, cust_ids)
                        if logs.empty:
                            st.warning("没有找到任何跟进记录")
                        else:
                            df_export = logs[['customer_name', 'customer_company', 'user_name', 'display_time', 'subject', 'content', 'participants', 'location', 'next_plan']]
                            df_export.columns = ['客户名称', '公司名称', '跟进人', '跟进时间', '主题', '内容', '相关人员', '地点', '下一步计划']
                            output = io.BytesIO()
                            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                df_export.to_excel(writer, sheet_name='全部跟进日志', index=False)
                            output.seek(0)
                            st.download_button(
                                label="📥 下载 Excel",
                                data=output,
                                file_name=f"全部跟进日志_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
        
        st.divider()
        
        # 查看跟进记录
        st.markdown("### 📋 查看跟进记录")
        cust_options_view = {f"{row['id']} - {row['name']} ({row['company']})": row['id'] for _, row in cust_list.iterrows()}
        selected_label_view = st.selectbox("选择客户", list(cust_options_view.keys()), key="view_cust_select")
        selected_cust_id_view = cust_options_view[selected_label_view]
        
        logs = query_df("""
            SELECT *, coalesce(log_time, created_at) as display_time
            FROM follow_logs
            WHERE ref_type='customer' AND ref_id=?
            ORDER BY display_time DESC
        """, (selected_cust_id_view,))
        
        if not logs.empty:
            user_map = get_user_map()
            logs["user_name"] = logs["user_id"].map(user_map).fillna(logs["user_id"])
            
            # 跟进记录显示
            for _, log_row in logs.iterrows():
                with st.expander(f"📝 {log_row['user_name']} - {pd.to_datetime(log_row['display_time']).strftime('%Y-%m-%d %H:%M')}", expanded=False):
                    cols = st.columns([4, 1])
                    with cols[0]:
                        if log_row["subject"]:
                            st.markdown(f"**主题**: {log_row['subject']}")
                        if log_row["participants"]:
                            st.markdown(f"**相关人员**: {log_row['participants']}")
                        if log_row["location"]:
                            st.markdown(f"**地点**: {log_row['location']}")
                        if log_row["next_plan"]:
                            st.markdown(f"**下一步计划**: {log_row['next_plan']}")
                        st.markdown(f"**内容**: {log_row['content']}")
                    
                    with cols[1]:
                        # 编辑和删除按钮
                        if is_boss or log_row["user_id"] == uid:
                            with st.popover("✏️ 编辑"):
                                with st.form(f"edit_follow_{log_row['id']}"):
                                    edit_content = st.text_area("跟进内容 *", value=log_row['content'])
                                    edit_subject = st.text_input("主题", value=log_row["subject"] if pd.notna(log_row["subject"]) else "")
                                    edit_participants = st.text_input("相关人员", value=log_row["participants"] if pd.notna(log_row["participants"]) else "")
                                    edit_location = st.text_input("地点", value=log_row["location"] if pd.notna(log_row["location"]) else "")
                                    edit_next_plan = st.text_area("下一步工作计划", value=log_row["next_plan"] if pd.notna(log_row["next_plan"]) else "")
                                    default_time = pd.to_datetime(log_row["log_time"]) if pd.notna(log_row["log_time"]) else pd.to_datetime(log_row["created_at"])
                                    edit_date = st.date_input("日期", value=default_time.date() if pd.notna(default_time) else date.today())
                                    edit_time = st.time_input("时间", value=default_time.time() if pd.notna(default_time) else datetime.now().time())
                                    if st.form_submit_button("保存修改"):
                                        with st.spinner("正在更新跟进记录..."):
                                            combined = datetime.combine(edit_date, edit_time)
                                            time_str = combined.strftime("%Y-%m-%d %H:%M:%S")
                                            update_sql = """
                                                UPDATE follow_logs
                                                SET content=?, subject=?, participants=?, location=?, next_plan=?, log_time=?
                                                WHERE id=?
                                            """
                                            try:
                                                execute_sql(update_sql, (edit_content, edit_subject, edit_participants, edit_location, edit_next_plan, time_str, log_row['id']))
                                                st.success("更新成功")
                                                time.sleep(0.5)
                                                st.rerun()
                                            except Exception as e:
                                                st.error(f"更新失败: {e}")
                        
                        if is_boss or log_row["user_id"] == uid:
                            with st.popover("🗑️ 删除"):
                                st.warning("确定删除此跟进记录？")
                                if st.button("确认删除", key=f"del_follow_{log_row['id']}", type="primary"):
                                    with st.spinner("正在删除跟进记录..."):
                                        try:
                                            execute_sql("DELETE FROM follow_logs WHERE id=?", (log_row['id'],))
                                            st.success("删除成功")
                                            time.sleep(0.5)
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"删除失败: {e}")
                        
                        if DOCX_AVAILABLE:
                            with st.popover("📄 导出Word"):
                                cust_info = cust_list[cust_list['id'] == selected_cust_id_view].iloc[0]
                                cust_name = cust_info['name']
                                cust_company = cust_info['company'] or ""
                                default_our_staff = st.session_state.get("u_info", {}).get("name", "")
                                default_visit_purpose = log_row["subject"] if pd.notna(log_row["subject"]) else ""
                                default_talk_record = log_row["content"]
                                default_followup = log_row["next_plan"] if pd.notna(log_row["next_plan"]) else ""
                                default_location = log_row["location"] if pd.notna(log_row["location"]) else ""
                                default_visit_date = pd.to_datetime(log_row["display_time"]).strftime("%Y-%m-%d") if pd.notna(log_row["display_time"]) else date.today().strftime("%Y-%m-%d")
                                
                                visit_date = st.text_input("拜访日期", value=default_visit_date, key=f"vd_{log_row['id']}")
                                location = st.text_input("地点", value=default_location, key=f"loc_{log_row['id']}")
                                our_staff = st.text_input("我方参会人员", value=default_our_staff, key=f"os_{log_row['id']}")
                                their_staff = st.text_input("对方参会人员", value="", key=f"ts_{log_row['id']}")
                                visit_purpose = st.text_input("拜访目的", value=default_visit_purpose, key=f"vp_{log_row['id']}")
                                talk_record = st.text_area("会谈记录", value=default_talk_record, height=150, key=f"tr_{log_row['id']}")
                                cooperation = st.text_area("可合作项目内容", value="", height=100, key=f"coop_{log_row['id']}")
                                followup_plan = st.text_area("后期跟进计划及建议", value=default_followup, height=80, key=f"fp_{log_row['id']}")
                                
                                if st.button("生成Word文档", key=f"gen_word_{log_row['id']}"):
                                    with st.spinner("正在生成Word文件..."):
                                        try:
                                            doc = Document()
                                            title = doc.add_heading('商 务 拜 访 记 录', level=1)
                                            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            data = [
                                                ("填表人", st.session_state.get("u_info", {}).get("name", "")),
                                                ("时间", visit_date),
                                                ("拜访单位", f"{cust_name}（{cust_company}）"),
                                                ("地点", location),
                                                ("我方参会人员", our_staff),
                                                ("对方参会人员", their_staff),
                                                ("拜访目的", visit_purpose),
                                                ("会谈记录（会谈要点及对方关注内容）", talk_record),
                                                ("可合作项目内容", cooperation),
                                                ("后期跟进计划及建议", followup_plan)
                                            ]
                                            table = doc.add_table(rows=len(data), cols=2)
                                            table.style = 'Table Grid'
                                            for i, (label, value) in enumerate(data):
                                                row = table.rows[i]
                                                row.cells[0].text = label
                                                row.cells[1].text = value
                                            for row in table.rows:
                                                for cell in row.cells:
                                                    if cell.paragraphs[0].runs:
                                                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                                            word_io = io.BytesIO()
                                            doc.save(word_io)
                                            word_io.seek(0)
                                            st.success("文档已生成，请点击下方下载按钮")
                                            st.download_button(
                                                label="📥 下载 Word 文档",
                                                data=word_io.getvalue(),
                                                file_name=f"拜访记录_{cust_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                key=f"dl_{log_row['id']}"
                                            )
                                        except Exception as e:
                                            st.error(f"生成Word失败: {e}")
                    
                    st.divider()
        else:
            st.info("暂无跟进记录")
        
        # 添加新跟进记录
        st.divider()
        st.markdown("### ✨ 添加新跟进记录")
        with st.form("add_follow", clear_on_submit=True):
            new_content = st.text_area("跟进内容 *")
            new_subject = st.text_input("主题")
            new_participants = st.text_input("相关人员")
            new_location = st.text_input("地点")
            new_next_plan = st.text_area("下一步工作计划")
            col_date, col_time = st.columns(2)
            with col_date:
                follow_date = st.date_input("跟进日期", value=None, key="cust_follow_date")
            with col_time:
                follow_time = st.time_input("跟进时间", value=None, step=60, key="cust_follow_time")
            
            if st.form_submit_button("提交跟进", type="primary"):
                if not new_content.strip():
                    st.error("内容不能为空")
                else:
                    with st.spinner("正在添加跟进..."):
                        actual_date = follow_date if follow_date else date.today()
                        actual_time = follow_time if follow_time else datetime.min.time()
                        combined = datetime.combine(actual_date, actual_time)
                        time_str = combined.strftime("%Y-%m-%d %H:%M:%S")
                        sql = """
                            INSERT INTO follow_logs
                            (ref_type, ref_id, user_id, content, subject, participants, location, next_plan, log_time)
                            VALUES ('customer', ?, ?, ?, ?, ?, ?, ?, ?)
                        """
                        try:
                            execute_sql(sql, (selected_cust_id_view, uid, new_content.strip(), new_subject, new_participants, new_location, new_next_plan, time_str))
                            from utils import update_customer_last_follow
                            update_customer_last_follow(selected_cust_id_view, combined)
                            st.success("跟进已添加")
                            time.sleep(0.5)
                            st.rerun()
                        except Exception as e:
                            st.error(f"添加失败: {e}")
